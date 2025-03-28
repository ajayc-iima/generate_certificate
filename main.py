from fastapi import FastAPI, File, UploadFile
from fastapi.responses import FileResponse, HTMLResponse
import pandas as pd
import os
import tempfile
import shutil
from docx import Document

app = FastAPI(title="IIMA Certificate Generator", description="Generate certificates efficiently using an Excel file and a Word template.")

@app.get("/", response_class=HTMLResponse)
async def serve_form():
    return """
   <!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Certificate Generator</title>
    <style>
        :root {
            --primary-color: #00457C;
            --secondary-color: #003366;
            --background-light: #f4f4f8;
            --text-color: #333;
            --success-color: #28a745;
            --error-color: #dc3545;
        }

        * {
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }

        .certificate-generator {
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Oxygen, Ubuntu, Cantarell, "Open Sans", "Helvetica Neue", sans-serif;
            max-width: 500px;
            margin: 0 auto;
            padding: 20px;
            background-color: white;
        }

        .generator-header {
            text-align: center;
            margin-bottom: 20px;
        }

        .generator-header h2 {
            color: var(--primary-color);
            font-weight: 600;
            margin-bottom: 10px;
        }

        .generator-header p {
            color: #666;
            font-size: 0.9rem;
        }

        .file-upload-container {
            background-color: var(--background-light);
            border-radius: 8px;
            padding: 20px;
            border: 2px dashed var(--primary-color);
            text-align: center;
            transition: all 0.3s ease;
        }

        .file-upload-container:hover {
            background-color: rgba(0, 69, 124, 0.05);
        }

        .file-input-wrapper {
            position: relative;
            margin-bottom: 15px;
        }

        .file-input-wrapper input[type="file"] {
            position: absolute;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            opacity: 0;
            cursor: pointer;
        }

        .file-input-label {
            display: inline-block;
            padding: 10px 20px;
            background-color: var(--primary-color);
            color: white;
            border-radius: 5px;
            cursor: pointer;
            transition: background-color 0.3s ease;
        }

        .file-input-label:hover {
            background-color: var(--secondary-color);
        }

        .file-name {
            margin-top: 10px;
            color: var(--text-color);
            font-size: 0.9rem;
        }

        .generate-btn {
            width: 100%;
            padding: 12px;
            background-color: var(--primary-color);
            color: white;
            border: none;
            border-radius: 5px;
            cursor: pointer;
            font-weight: 600;
            transition: background-color 0.3s ease;
        }

        .generate-btn:hover {
            background-color: var(--secondary-color);
        }

        .generate-btn:disabled {
            background-color: #cccccc;
            cursor: not-allowed;
        }

        .status-message {
            margin-top: 15px;
            text-align: center;
            font-size: 0.9rem;
        }

        .status-message.success {
            color: var(--success-color);
        }

        .status-message.error {
            color: var(--error-color);
        }

        .progress-container {
            width: 100%;
            background-color: #e0e0e0;
            border-radius: 5px;
            margin-top: 15px;
            display: none;
        }

        .progress-bar {
            width: 0;
            height: 5px;
            background-color: var(--success-color);
            border-radius: 5px;
            transition: width 0.5s ease;
        }

        .download-link {
            display: block;
            margin-top: 15px;
            text-align: center;
            padding: 10px;
            background-color: var(--success-color);
            color: white;
            text-decoration: none;
            border-radius: 5px;
        }
    </style>
</head>
<body>
    <div class="certificate-generator">
        <div class="generator-header">
            <h2>Certificate Generator</h2>
            <p>Upload an Excel file with participant names and a Word template</p>
        </div>

        <form id="certificate-form">
            <div class="file-upload-container">
                <div class="file-input-wrapper">
                    <input 
                        type="file" 
                        id="excel-file" 
                        name="excel" 
                        accept=".xlsx,.xls" 
                        required
                    >
                    <label for="excel-file" class="file-input-label">
                        Select Excel File
                    </label>
                    <div id="excel-file-name" class="file-name">No file selected</div>
                </div>

                <div class="file-input-wrapper">
                    <input 
                        type="file" 
                        id="docx-file" 
                        name="docx" 
                        accept=".docx" 
                        required
                    >
                    <label for="docx-file" class="file-input-label">
                        Select Word Template
                    </label>
                    <div id="docx-file-name" class="file-name">No file selected</div>
                </div>

                <div class="progress-container">
                    <div class="progress-bar" id="progress-bar"></div>
                </div>

                <button 
                    type="submit" 
                    class="generate-btn" 
                    id="generate-btn"
                    disabled
                >
                    Generate Certificates
                </button>
            </div>
        </form>

        <div id="status-message" class="status-message"></div>
        <a 
            href="#" 
            id="download-link" 
            class="download-link" 
            style="display:none;"
            download
        >
            Download Certificates
        </a>
    </div>

    <script>
        document.addEventListener('DOMContentLoaded', () => {
            const form = document.getElementById('certificate-form');
            const excelInput = document.getElementById('excel-file');
            const docxInput = document.getElementById('docx-file');
            const excelFileName = document.getElementById('excel-file-name');
            const docxFileName = document.getElementById('docx-file-name');
            const generateBtn = document.getElementById('generate-btn');
            const statusMessage = document.getElementById('status-message');
            const progressContainer = document.querySelector('.progress-container');
            const progressBar = document.getElementById('progress-bar');
            const downloadLink = document.getElementById('download-link');

            // File input event listeners
            [excelInput, docxInput].forEach(input => {
                input.addEventListener('change', (e) => {
                    const fileName = e.target.files[0] ? e.target.files[0].name : 'No file selected';
                    const fileNameEl = e.target.id === 'excel-file' ? excelFileName : docxFileName;
                    fileNameEl.textContent = fileName;

                    // Enable generate button only when both files are selected
                    generateBtn.disabled = !(excelInput.files.length && docxInput.files.length);
                });
            });

            form.addEventListener('submit', async (e) => {
                e.preventDefault();

                // Reset UI
                statusMessage.textContent = '';
                statusMessage.className = 'status-message';
                progressContainer.style.display = 'none';
                progressBar.style.width = '0%';
                downloadLink.style.display = 'none';

                const excelFile = excelInput.files[0];
                const docxFile = docxInput.files[0];

                // Validate file sizes
                const MAX_EXCEL_SIZE = 10 * 1024 * 1024; // 10MB
                const MAX_DOCX_SIZE = 5 * 1024 * 1024;   // 5MB

                if (excelFile.size > MAX_EXCEL_SIZE) {
                    statusMessage.textContent = 'Excel file is too large (max 10MB)';
                    statusMessage.classList.add('error');
                    return;
                }

                if (docxFile.size > MAX_DOCX_SIZE) {
                    statusMessage.textContent = 'Word template is too large (max 5MB)';
                    statusMessage.classList.add('error');
                    return;
                }

                const formData = new FormData();
                formData.append('excel_file', excelFile);
                formData.append('docx_template', docxFile);

                try {
                    // Disable button during processing
                    generateBtn.disabled = true;
                    generateBtn.textContent = 'Generating...';

                    // Show progress
                    progressContainer.style.display = 'block';

                    const response = await fetch('/generate-certificates/', {
                        method: 'POST',
                        body: formData,
                    });

                    if (!response.ok) {
                        throw new Error(await response.text() || 'Certificate generation failed');
                    }

                    const blob = await response.blob();
                    const url = window.URL.createObjectURL(blob);
                    
                    downloadLink.href = url;
                    downloadLink.download = `certificates_${new Date().toISOString().split('T')[0]}.zip`;
                    
                    statusMessage.textContent = 'Certificates generated successfully!';
                    statusMessage.classList.add('success');
                    downloadLink.style.display = 'block';

                } catch (error) {
                    console.error('Generation error:', error);
                    statusMessage.textContent = error.message || 'An unexpected error occurred';
                    statusMessage.classList.add('error');
                } finally {
                    // Reset button
                    generateBtn.disabled = false;
                    generateBtn.textContent = 'Generate Certificates';
                    progressContainer.style.display = 'none';
                }
            });
        });
    </script>
</body>
</html>
    """

@app.post("/generate-certificates/")
async def generate_certificates(excel_file: UploadFile = File(...), docx_template: UploadFile = File(...)):
    temp_dir = tempfile.mkdtemp()
    output_dir = os.path.join(temp_dir, "certificates")
    os.makedirs(output_dir, exist_ok=True)

    # Save uploaded files
    excel_path = os.path.join(temp_dir, excel_file.filename)
    docx_path = os.path.join(temp_dir, docx_template.filename)
    with open(excel_path, "wb") as f:
        shutil.copyfileobj(excel_file.file, f)
    with open(docx_path, "wb") as f:
        shutil.copyfileobj(docx_template.file, f)

    # Read Excel file
    df = pd.read_excel(excel_path)
    names = df.iloc[:, 0].dropna().str.strip().tolist()

    for name in names:
        doc = Document(docx_path)
        replace_text_preserving_format(doc, "{Name}", name)  # ✅ Format-safe replacement
        doc.save(os.path.join(output_dir, f"{name}_Certificate.docx"))

    # Zip all certificates
    zip_path = os.path.join(temp_dir, "certificates.zip")
    shutil.make_archive(zip_path.replace(".zip", ""), 'zip', output_dir)

    return FileResponse(zip_path, filename="certificates.zip", media_type="application/zip")

def replace_text_preserving_format(doc, placeholder, replacement):
    """ Replaces {Name} without losing formatting & color """
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)

    # Also replace inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement)
